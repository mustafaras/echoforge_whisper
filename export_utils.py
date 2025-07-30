"""
Export Utilities for Whisper AI
Bu modül PDF, Word, Excel, QR kod ve ZIP export işlemlerini yönetir.
"""

import streamlit as st
import os
import io
import time
from datetime import datetime
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

# =============================================
# 🌐 TRANSLATION UTILITIES
# =============================================

class TranslationHelper:
    """Çeviri yardımcı sınıfı"""
    
    @staticmethod
    def translate_text(text: str, target_language: str = "en", openai_client=None) -> str:
        """OpenAI API kullanarak metin çevirisi - Gelişmiş error handling ile"""
        
        # Debug logging
        logger.info(f"Translation request: target={target_language}, length={len(text)}")
        
        try:
            if not openai_client:
                logger.error("OpenAI client not provided for translation")
                raise Exception("OpenAI client bulunamadı")
            
            if not text or len(text.strip()) == 0:
                logger.warning("Empty text provided for translation")
                return text
            
            # Text uzunluk kontrolü
            if len(text) > 15000:  # Çok uzun metinleri parçala
                logger.info(f"Text too long ({len(text)} chars), splitting into chunks")
                max_chunk_size = 4000
                chunks = [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]
                translated_chunks = []
                
                for i, chunk in enumerate(chunks):
                    logger.info(f"Translating chunk {i+1}/{len(chunks)}")
                    try:
                        translated_chunk = TranslationHelper.translate_text(chunk, target_language, openai_client)
                        translated_chunks.append(translated_chunk)
                        time.sleep(0.5)  # API rate limit
                    except Exception as chunk_error:
                        logger.error(f"Chunk {i+1} translation failed: {chunk_error}")
                        translated_chunks.append(chunk)  # Hata durumunda orijinal chunk
                
                result = " ".join(translated_chunks)
                logger.info(f"Chunked translation completed, result length: {len(result)}")
                return result
            
            # Dil kodlarını belirle
            lang_prompts = {
                "en": "Translate the following Turkish text to English. Preserve formatting and provide only the translation without any explanations:",
                "tr": "Translate the following English text to Turkish. Preserve formatting and provide only the translation without any explanations:",
                "auto": "Detect the language and translate to the opposite (Turkish↔English). Preserve formatting and provide only the translation:"
            }
            
            # Çeviri promptu
            prompt = lang_prompts.get(target_language, lang_prompts["en"])
            
            logger.info(f"Starting OpenAI translation request for {target_language}")
            
            # OpenAI API çağrısı - daha sağlam parametreler
            response = openai_client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[
                    {
                        "role": "system", 
                        "content": "You are a professional translator. Provide accurate, contextual translations while preserving formatting, tone, and meaning. Only return the translated text, no explanations or additional content."
                    },
                    {
                        "role": "user", 
                        "content": f"{prompt}\n\nText to translate:\n{text}"
                    }
                ],
                max_tokens=4000,
                temperature=0.1,
                top_p=1.0,
                frequency_penalty=0.0,
                presence_penalty=0.0
            )
            
            if not response or not response.choices:
                raise Exception("OpenAI API boş yanıt döndü")
            
            translated_text = response.choices[0].message.content.strip()
            
            if not translated_text:
                raise Exception("Çeviri sonucu boş")
            
            # Sonuç validasyonu
            if len(translated_text) < len(text) * 0.3:  # Çok kısa sonuç
                logger.warning(f"Translation result seems too short: {len(translated_text)} vs {len(text)}")
            
            logger.info(f"Translation completed successfully: {len(text)} -> {len(translated_text)} chars")
            return translated_text
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Translation failed: {error_msg}")
            
            # Specific error handling
            if "rate_limit" in error_msg.lower():
                st.error("❌ API rate limit aşıldı. Lütfen biraz bekleyip tekrar deneyin.")
            elif "authentication" in error_msg.lower():
                st.error("❌ OpenAI API anahtarı geçersiz.")
            elif "insufficient_quota" in error_msg.lower():
                st.error("❌ OpenAI API quota'sı dolmuş.")
            else:
                st.error(f"❌ Çeviri hatası: {error_msg}")
            
            # Hata durumunda orijinal metni döndür
            return text
    
    @staticmethod
    def detect_language(text: str) -> str:
        """Basit dil tespiti (Türkçe/İngilizce)"""
        try:
            # Türkçe karakterlerin varlığını kontrol et
            turkish_chars = ['ç', 'ğ', 'ı', 'ş', 'ü', 'ö', 'Ç', 'Ğ', 'İ', 'Ş', 'Ü', 'Ö']
            
            # Türkçe kelimeler (yaygın olanlar)
            turkish_words = ['bir', 'bu', 'şu', 'o', 'ben', 'sen', 'biz', 'siz', 'onlar', 'var', 'yok', 'için', 'ile', 'kadar', 'gibi', 'çok', 'az', 'büyük', 'küçük']
            
            text_lower = text.lower()
            
            # Türkçe karakter kontrolü
            turkish_char_count = sum(1 for char in text if char in turkish_chars)
            if turkish_char_count > 0:
                return "tr"
            
            # Türkçe kelime kontrolü
            turkish_word_count = sum(1 for word in turkish_words if word in text_lower)
            total_words = len(text_lower.split())
            
            if total_words > 0 and (turkish_word_count / total_words) > 0.1:
                return "tr"
            
            return "en"  # Varsayılan olarak İngilizce
            
        except:
            return "auto"

class PDFExporter:
    """Gelişmiş PDF export sınıfı"""
    
    @staticmethod
    def create_translated_pdf_report(transcription_data, ai_analysis=None, target_language="en", openai_client=None):
        """Çevirili PDF raporu oluştur - Gelişmiş hata yönetimi ile"""
        
        # Debug logging
        logger.info(f"Starting translated PDF creation for language: {target_language}")
        
        try:
            if not openai_client:
                st.error("❌ OpenAI client bulunamadı!")
                logger.error("OpenAI client is None")
                return None
                
            # Çeviri helper'ı kullan
            translator = TranslationHelper()
            
            # Progress tracking container
            progress_placeholder = st.empty()
            
            with progress_placeholder.container():
                # Progress bar ve status
                progress_bar = st.progress(0)
                status_text = st.empty()
            
                status_text.info("🔄 Çeviri işlemi başlatılıyor...")
                progress_bar.progress(5)
                
                # Ana metni kontrol et
                original_text = transcription_data.get('transcript_text', '')
                
                if not original_text:
                    st.error("❌ Çevrilecek metin bulunamadı!")
                    logger.error("No transcript text found")
                    return None
                
                if len(original_text.strip()) == 0:
                    st.error("❌ Transkripsiyon metni boş!")
                    logger.error("Transcript text is empty")
                    return None
                
                # Metin uzunluğu bilgisi
                char_count = len(original_text)
                word_count = len(original_text.split())
                
                status_text.info(f"📝 Çevrilecek metin: {word_count} kelime, {char_count} karakter")
                progress_bar.progress(10)
                
                # Ana metin çevirisi
                logger.info(f"Starting main text translation, length: {char_count}")
                status_text.info(f"🌐 Ana metin çevriliyor... ({target_language.upper()})")
                
                try:
                    translated_text = translator.translate_text(original_text, target_language, openai_client)
                    
                    if not translated_text or translated_text.strip() == "":
                        raise Exception("Çeviri sonucu boş döndü")
                    
                    # Çeviri başarılı mı kontrol et
                    if translated_text == original_text:
                        st.warning("⚠️ Çeviri yapılamadı, orijinal metin kullanılacak")
                        logger.warning("Translation returned same text as original")
                    else:
                        translated_word_count = len(translated_text.split())
                        st.success(f"✅ Ana metin çevrildi! ({translated_word_count} kelime)")
                        logger.info(f"Main text translation successful, result length: {len(translated_text)}")
                    
                except Exception as translation_error:
                    logger.error(f"Main text translation failed: {translation_error}")
                    st.error(f"❌ Ana metin çevirisi başarısız: {str(translation_error)}")
                    st.info("🔄 Orijinal metin ile devam ediliyor...")
                    translated_text = original_text
                
                progress_bar.progress(50)
                
                # AI analiz çevirisi (varsa)
                translated_analysis = None
                if ai_analysis:
                    status_text.info("🤖 AI analiz verileri çevriliyor...")
                    translated_analysis = {}
                    
                    # Özet çevirisi
                    if ai_analysis.get('summary'):
                        try:
                            logger.info("Translating summary")
                            summary_translation = translator.translate_text(
                                ai_analysis['summary'], target_language, openai_client
                            )
                            translated_analysis['summary'] = summary_translation
                            st.success("✅ Özet çevrildi")
                        except Exception as e:
                            logger.error(f"Summary translation failed: {e}")
                            translated_analysis['summary'] = ai_analysis['summary']
                            st.warning("⚠️ Özet çevirisi başarısız, orijinal kullanılıyor")
                    
                    # Anahtar kelime çevirisi
                    if ai_analysis.get('keywords'):
                        try:
                            logger.info("Translating keywords")
                            translated_keywords = []
                            keywords_to_translate = ai_analysis['keywords'][:5]  # İlk 5'i çevir
                            
                            for i, keyword in enumerate(keywords_to_translate):
                                try:
                                    translated_keyword = translator.translate_text(
                                        keyword, target_language, openai_client
                                    )
                                    translated_keywords.append(translated_keyword)
                                    time.sleep(0.2)  # API rate limit
                                except:
                                    translated_keywords.append(keyword)  # Hata durumunda orijinal
                            
                            # Kalan anahtar kelimeleri orijinal olarak ekle
                            if len(ai_analysis['keywords']) > 5:
                                translated_keywords.extend(ai_analysis['keywords'][5:])
                                
                            translated_analysis['keywords'] = translated_keywords
                            st.success("✅ Anahtar kelimeler çevrildi")
                            
                        except Exception as e:
                            logger.error(f"Keywords translation failed: {e}")
                            translated_analysis['keywords'] = ai_analysis['keywords']
                            st.warning("⚠️ Anahtar kelime çevirisi başarısız, orijinal kullanılıyor")
                    
                    # Diğer analiz verilerini kopyala
                    for key in ['emotion_analysis', 'speed_analysis']:
                        if ai_analysis.get(key):
                            translated_analysis[key] = ai_analysis[key]
                            
                    progress_bar.progress(75)
                
                # PDF oluşturma
                status_text.info("📄 Çevirili PDF oluşturuluyor...")
                
                # Transcription data'yı güncelle
                translated_transcription_data = transcription_data.copy()
                translated_transcription_data['transcript_text'] = translated_text
                translated_transcription_data['original_text'] = original_text
                translated_transcription_data['translation_language'] = target_language
                
                # Dil bilgisini ekle
                lang_name = "İngilizce" if target_language == "en" else "Türkçe"
                translated_transcription_data['translation_info'] = {
                    'target_language': target_language,
                    'language_name': lang_name,
                    'original_length': len(original_text),
                    'translated_length': len(translated_text),
                    'translation_time': time.strftime('%Y-%m-%d %H:%M:%S')
                }
                
                logger.info("Creating PDF with translated content")
                
                # PDF oluştur
                pdf_data = PDFExporter.create_advanced_pdf_report(
                    translated_transcription_data, 
                    translated_analysis
                )
                
                if pdf_data:
                    progress_bar.progress(100)
                    status_text.success(f"✅ {lang_name} PDF başarıyla oluşturuldu!")
                    logger.info(f"Translated PDF creation completed successfully")
                    
                    # Progress container'ı temizle
                    time.sleep(1.5)
                    progress_placeholder.empty()
                    
                    return pdf_data
                else:
                    st.error("❌ PDF oluşturma başarısız")
                    logger.error("PDF creation returned None")
                    return None
            
        except Exception as e:
            logger.error(f"Translated PDF creation failed: {e}")
            st.error(f"❌ Çevirili PDF oluşturma hatası: {str(e)}")
            
            # Progress container'ı temizle
            if 'progress_placeholder' in locals():
                progress_placeholder.empty()
            
            return None
    
    @staticmethod
    def create_advanced_pdf_report(transcription_data, ai_analysis=None):
        """Gelişmiş PDF raporu oluştur - Türkçe karakter desteği ile"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
            from io import BytesIO
            import os
            from datetime import datetime
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            story = []
            
            # Türkçe font desteği için sistemdeki fontları kullan
            font_name = 'Helvetica'
            font_bold = 'Helvetica-Bold'
            
            # Önce Windows sistem fontlarını dene
            font_paths = [
                # Arial fontları (en yaygın)
                ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/arialbd.ttf', 'Arial'),
                ('C:/Windows/Fonts/calibri.ttf', 'C:/Windows/Fonts/calibrib.ttf', 'Calibri'),
                ('C:/Windows/Fonts/tahoma.ttf', 'C:/Windows/Fonts/tahomabd.ttf', 'Tahoma'),
                ('C:/Windows/Fonts/verdana.ttf', 'C:/Windows/Fonts/verdanab.ttf', 'Verdana'),
                # DejaVu fontları
                ('C:/Windows/Fonts/DejaVuSans.ttf', 'C:/Windows/Fonts/DejaVuSans-Bold.ttf', 'DejaVuSans'),
                # Alternatif yollar
                ('arial.ttf', 'arialbd.ttf', 'Arial'),
                ('calibri.ttf', 'calibrib.ttf', 'Calibri'),
            ]
            
            font_loaded = False
            for regular_path, bold_path, font_family in font_paths:
                try:
                    import os
                    if os.path.exists(regular_path) and os.path.exists(bold_path):
                        pdfmetrics.registerFont(TTFont(font_family, regular_path))
                        pdfmetrics.registerFont(TTFont(f'{font_family}-Bold', bold_path))
                        font_name = font_family
                        font_bold = f'{font_family}-Bold'
                        font_loaded = True
                        break
                except Exception as e:
                    continue
            
            # Eğer hiç font yüklenemezse, UTF-8 karakterleri destekleyen bir font kullan
            if not font_loaded:
                try:
                    # Son çare olarak reportlab ile gelen fontları kullan
                    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                    pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
                    font_name = 'HeiseiMin-W3'
                    font_bold = 'HeiseiMin-W3'
                    font_loaded = True
                except:
                    # Çok son çare olarak Helvetica kullan ama UTF-8 sorunları olabilir
                    font_name = 'Helvetica'
                    font_bold = 'Helvetica-Bold'
            
            # Özel stiller tanımla
            styles = getSampleStyleSheet()
            
            # Başlık stili
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontName=font_bold,
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#2E86AB')
            )
            
            # Alt başlık stili
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading1'],
                fontName=font_bold,
                fontSize=16,
                spaceAfter=12,
                spaceBefore=20,
                textColor=colors.HexColor('#A23B72'),
                borderWidth=0,
                borderColor=colors.HexColor('#A23B72'),
                borderPadding=5
            )
            
            # Normal metin stili
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                spaceAfter=6,
                alignment=TA_JUSTIFY,
                leading=14
            )
            
            # Bilgi kutusu stili
            info_style = ParagraphStyle(
                'InfoStyle',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=6,
                leftIndent=20,
                borderWidth=1,
                borderColor=colors.HexColor('#E8F4FD'),
                borderPadding=10,
                backColor=colors.HexColor('#F8FBFF')
            )
            
            # Ana başlık
            title_text = "Whisper AI Transkripsiyon Raporu"
            if font_loaded:
                title_text = "🎵 Whisper AI Transkripsiyon Raporu"
            
            title = Paragraph(title_text, title_style)
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Rapor bilgileri tablosu
            current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            
            # Çeviri bilgilerini kontrol et
            is_translated = transcription_data.get('translation_language') is not None
            original_text_exists = transcription_data.get('original_text') is not None
            
            if font_loaded:
                report_data = [
                    ['📁 Dosya Adı:', str(transcription_data.get('file_name', 'Bilinmiyor'))],
                    ['📅 İşlem Tarihi:', current_time],
                    ['🌍 Dil:', str(transcription_data.get('language', 'Otomatik'))],
                    ['⏱️ İşlem Süresi:', f"{transcription_data.get('processing_time', 0):.1f} saniye"],
                    ['📊 Dosya Boyutu:', f"{transcription_data.get('file_size_mb', 0):.1f} MB"],
                    ['🔢 ID:', str(transcription_data.get('id', 'N/A'))]
                ]
                
                # Çeviri bilgilerini ekle
                if is_translated:
                    translation_lang = transcription_data.get('translation_language')
                    lang_name = "İngilizce" if translation_lang == "en" else "Türkçe" if translation_lang == "tr" else translation_lang
                    report_data.append(['🌐 Çeviri Dili:', lang_name])
                    report_data.append(['🔄 Çeviri Durumu:', '✅ Çevirildi'])
                
            else:
                report_data = [
                    ['Dosya Adi:', str(transcription_data.get('file_name', 'Bilinmiyor'))],
                    ['Islem Tarihi:', current_time],
                    ['Dil:', str(transcription_data.get('language', 'Otomatik'))],
                    ['Islem Suresi:', f"{transcription_data.get('processing_time', 0):.1f} saniye"],
                    ['Dosya Boyutu:', f"{transcription_data.get('file_size_mb', 0):.1f} MB"],
                    ['ID:', str(transcription_data.get('id', 'N/A'))]
                ]
                
                # Çeviri bilgilerini ekle
                if is_translated:
                    translation_lang = transcription_data.get('translation_language')
                    lang_name = "Ingilizce" if translation_lang == "en" else "Turkce" if translation_lang == "tr" else translation_lang
                    report_data.append(['Ceviri Dili:', lang_name])
                    report_data.append(['Ceviri Durumu:', 'Cevirildi'])
            
            # Tablo stili
            report_table = Table(report_data, colWidths=[4*cm, 10*cm])
            report_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E8F4FD')),
                ('BACKGROUND', (1, 0), (1, -1), colors.white),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), font_bold),
                ('FONTNAME', (1, 0), (1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            
            story.append(report_table)
            story.append(Spacer(1, 25))
            
            # Transkripsiyon bölümü
            transcript_title_text = "Transkripsiyon Metni"
            if is_translated:
                translation_lang = transcription_data.get('translation_language')
                lang_name = "İngilizce" if translation_lang == "en" else "Türkçe"
                transcript_title_text = f"Çevirili Metin ({lang_name})"
            
            if font_loaded:
                if is_translated:
                    transcript_title_text = f"🌐 {transcript_title_text}"
                else:
                    transcript_title_text = "📝 Transkripsiyon Metni"
            
            transcript_title = Paragraph(transcript_title_text, subtitle_style)
            story.append(transcript_title)
            
            # Transkripsiyon metni
            transcript_text = transcription_data.get('transcript_text', 'Transkripsiyon metni bulunamadı.')
            
            # Türkçe karakter düzeltmesi (font yoksa)
            if not font_loaded:
                char_map = {
                    'ç': 'c', 'Ç': 'C', 'ğ': 'g', 'Ğ': 'G', 'ı': 'i', 'İ': 'I',
                    'ö': 'o', 'Ö': 'O', 'ş': 's', 'Ş': 'S', 'ü': 'u', 'Ü': 'U'
                }
                for tr_char, en_char in char_map.items():
                    transcript_text = transcript_text.replace(tr_char, en_char)
            
            # Metin istatistikleri
            word_count = len(transcript_text.split())
            char_count = len(transcript_text)
            
            stats_text = f"Metin Istatistikleri: {word_count} kelime, {char_count} karakter"
            if font_loaded:
                stats_text = f"📊 <b>Metin İstatistikleri:</b> {word_count} kelime, {char_count} karakter"
            
            stats_para = Paragraph(stats_text, info_style)
            story.append(stats_para)
            story.append(Spacer(1, 12))
            
            # Uzun metni parçalara böl
            max_chars_per_paragraph = 2000
            if len(transcript_text) > max_chars_per_paragraph:
                paragraphs = transcript_text.split('\n\n')
                current_paragraph = ""
                
                for para in paragraphs:
                    if len(current_paragraph + para) < max_chars_per_paragraph:
                        current_paragraph += para + "\n\n"
                    else:
                        if current_paragraph:
                            story.append(Paragraph(current_paragraph.strip(), normal_style))
                            story.append(Spacer(1, 8))
                        current_paragraph = para + "\n\n"
                
                if current_paragraph:
                    story.append(Paragraph(current_paragraph.strip(), normal_style))
            else:
                transcript_para = Paragraph(transcript_text, normal_style)
                story.append(transcript_para)
            
            # Orijinal metin bölümü (eğer çeviri varsa)
            if is_translated and original_text_exists:
                story.append(Spacer(1, 25))
                
                original_title_text = "Orijinal Metin"
                if font_loaded:
                    original_title_text = "📄 Orijinal Metin"
                
                original_title = Paragraph(original_title_text, subtitle_style)
                story.append(original_title)
                
                # Orijinal metin istatistikleri
                original_text = transcription_data.get('original_text', '')
                original_word_count = len(original_text.split())
                original_char_count = len(original_text)
                
                original_stats_text = f"Orijinal Metin Istatistikleri: {original_word_count} kelime, {original_char_count} karakter"
                if font_loaded:
                    original_stats_text = f"📊 <b>Orijinal Metin İstatistikleri:</b> {original_word_count} kelime, {original_char_count} karakter"
                
                original_stats_para = Paragraph(original_stats_text, info_style)
                story.append(original_stats_para)
                story.append(Spacer(1, 12))
                
                # Orijinal metni paragraf paragraf ekle
                if original_text:
                    # Türkçe karakter düzeltmesi (font yoksa)
                    original_display_text = original_text
                    if not font_loaded:
                        char_map = {
                            'ç': 'c', 'Ç': 'C', 'ğ': 'g', 'Ğ': 'G', 'ı': 'i', 'İ': 'I',
                            'ö': 'o', 'Ö': 'O', 'ş': 's', 'Ş': 'S', 'ü': 'u', 'Ü': 'U'
                        }
                        for tr_char, en_char in char_map.items():
                            original_display_text = original_display_text.replace(tr_char, en_char)
                    
                    # Uzun metinleri parçalara böl
                    if len(original_display_text) > 3000:
                        paragraphs = original_display_text.split('\n\n')
                        current_paragraph = ""
                        
                        for para in paragraphs:
                            if len(current_paragraph + para) > 2000:
                                if current_paragraph:
                                    story.append(Paragraph(current_paragraph.strip(), normal_style))
                                    story.append(Spacer(1, 8))
                                current_paragraph = para + "\n\n"
                            else:
                                current_paragraph += para + "\n\n"
                        
                        if current_paragraph:
                            story.append(Paragraph(current_paragraph.strip(), normal_style))
                    else:
                        original_para = Paragraph(original_display_text, normal_style)
                        story.append(original_para)
            
            # AI analiz bölümü (kapsamlı)
            if ai_analysis and ai_analysis.get('summary'):
                story.append(Spacer(1, 25))
                
                ai_title_text = "AI Analiz Sonuclari - Kapsamli Rapor"
                if font_loaded:
                    ai_title_text = "🤖 AI Analiz Sonuçları - Kapsamlı Rapor"
                
                ai_title = Paragraph(ai_title_text, subtitle_style)
                story.append(ai_title)
                
                # AI bilgileri
                ai_info = f"<b>Analiz Turu:</b> Gelismis Metin Analizi<br/><b>Model:</b> {ai_analysis.get('model', 'GPT-4-Turbo')}<br/><b>Analiz Tarihi:</b> {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
                if font_loaded:
                    ai_info = f"🧠 <b>Analiz Türü:</b> Gelişmiş Metin Analizi<br/>⚙️ <b>Model:</b> {ai_analysis.get('model', 'GPT-4-Turbo')}<br/>📊 <b>Analiz Tarihi:</b> {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
                
                ai_info_para = Paragraph(ai_info, info_style)
                story.append(ai_info_para)
                story.append(Spacer(1, 15))
                
                # Özet
                if ai_analysis.get('summary'):
                    summary_text_content = str(ai_analysis.get('summary', ''))
                    if not font_loaded:
                        char_map = {
                            'ç': 'c', 'Ç': 'C', 'ğ': 'g', 'Ğ': 'G', 'ı': 'i', 'İ': 'I',
                            'ö': 'o', 'Ö': 'O', 'ş': 's', 'Ş': 'S', 'ü': 'u', 'Ü': 'U'
                        }
                        for tr_char, en_char in char_map.items():
                            summary_text_content = summary_text_content.replace(tr_char, en_char)
                    
                    summary_text = f"<b>Metin Ozeti:</b><br/><br/>{summary_text_content}"
                    if font_loaded:
                        summary_text = f"<b>📋 Metin Özeti:</b><br/><br/>{summary_text_content}"
                    
                    summary_para = Paragraph(summary_text, normal_style)
                    story.append(summary_para)
                    story.append(Spacer(1, 15))
                
                # Anahtar kelimeler
                if ai_analysis.get('keywords'):
                    keywords_list = ai_analysis['keywords']
                    if not font_loaded:
                        char_map = {
                            'ç': 'c', 'Ç': 'C', 'ğ': 'g', 'Ğ': 'G', 'ı': 'i', 'İ': 'I',
                            'ö': 'o', 'Ö': 'O', 'ş': 's', 'Ş': 'S', 'ü': 'u', 'Ü': 'U'
                        }
                        keywords_list = [
                            ''.join(char_map.get(c, c) for c in keyword) 
                            for keyword in keywords_list
                        ]
                    
                    keywords_title = "<b>Anahtar Kelimeler:</b>"
                    if font_loaded:
                        keywords_title = "🔑 <b>Anahtar Kelimeler:</b>"
                    
                    keywords_title_para = Paragraph(keywords_title, normal_style)
                    story.append(keywords_title_para)
                    
                    # Anahtar kelimeleri tablo formatında göster
                    keywords_chunks = [keywords_list[i:i+3] for i in range(0, len(keywords_list), 3)]
                    keywords_table_data = []
                    for chunk in keywords_chunks:
                        row = chunk + [''] * (3 - len(chunk))
                        keywords_table_data.append(row)
                    
                    if keywords_table_data:
                        keywords_table = Table(keywords_table_data, colWidths=[5*cm, 5*cm, 5*cm])
                        keywords_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F0F8FF')),
                            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, -1), font_name),
                            ('FONTSIZE', (0, 0), (-1, -1), 10),
                            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#DDDDDD')),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('LEFTPADDING', (0, 0), (-1, -1), 8),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ]))
                        story.append(keywords_table)
                        story.append(Spacer(1, 15))
            
            # Footer
            story.append(Spacer(1, 30))
            footer_text = f"<i>Bu rapor Whisper AI tarafindan {current_time} tarihinde otomatik olarak olusturulmustur.</i>"
            if font_loaded:
                footer_text = f"<i>Bu rapor Whisper AI tarafından {current_time} tarihinde otomatik olarak oluşturulmuştur.</i>"
            
            footer_para = Paragraph(footer_text, ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=8,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#666666')
            ))
            story.append(footer_para)
            
            # PDF'i oluştur
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"PDF creation error: {e}")
            st.error(f"PDF oluşturma hatası: {str(e)}")
            return None


class WordExporter:
    """Gelişmiş Word export sınıfı"""
    
    @staticmethod
    def create_advanced_word_document(transcription_data, ai_analysis=None):
        """Gelişmiş Word belgesi oluştur"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.shared import OxmlElement, qn
            from io import BytesIO
            
            doc = Document()
            
            # Başlık
            title = doc.add_heading('Whisper AI Transkripsiyon Raporu', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Rapor bilgileri tablosu
            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_data = [
                ('Dosya Adı:', str(transcription_data.get('file_name', 'Bilinmiyor'))),
                ('İşlem Tarihi:', datetime.now().strftime("%d.%m.%Y %H:%M:%S")),
                ('Dil:', str(transcription_data.get('language', 'Otomatik'))),
                ('İşlem Süresi:', f"{transcription_data.get('processing_time', 0):.1f} saniye"),
                ('Dosya Boyutu:', f"{transcription_data.get('file_size_mb', 0):.1f} MB"),
                ('ID:', str(transcription_data.get('id', 'N/A')))
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = value
                info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
            
            doc.add_page_break()
            
            # Transkripsiyon bölümü
            doc.add_heading('Transkripsiyon Metni', level=1)
            
            transcript_text = transcription_data.get('transcript_text', 'Transkripsiyon metni bulunamadı.')
            
            # Metin istatistikleri
            word_count = len(transcript_text.split())
            char_count = len(transcript_text)
            
            stats_para = doc.add_paragraph()
            stats_para.add_run(f"Metin İstatistikleri: {word_count} kelime, {char_count} karakter").italic = True
            
            doc.add_paragraph()
            
            # Transkripsiyon metni - uzun metinleri parçalara böl
            max_chars_per_paragraph = 3000
            if len(transcript_text) > max_chars_per_paragraph:
                paragraphs = transcript_text.split('\n\n')
                current_paragraph = ""
                
                for para in paragraphs:
                    if len(current_paragraph + para) < max_chars_per_paragraph:
                        current_paragraph += para + "\n\n"
                    else:
                        if current_paragraph:
                            doc.add_paragraph(current_paragraph.strip())
                        current_paragraph = para + "\n\n"
                
                if current_paragraph:
                    doc.add_paragraph(current_paragraph.strip())
            else:
                doc.add_paragraph(transcript_text)
            
            # AI analiz bölümü
            if ai_analysis and ai_analysis.get('summary'):
                doc.add_page_break()
                doc.add_heading('AI Analiz Sonuçları', level=1)
                
                # AI bilgileri
                ai_info_para = doc.add_paragraph()
                ai_info_para.add_run(f"Analiz Türü: Gelişmiş Metin Analizi\n").bold = True
                ai_info_para.add_run(f"Model: {ai_analysis.get('model', 'GPT-4-Turbo')}\n")
                ai_info_para.add_run(f"Analiz Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
                
                doc.add_paragraph()
                
                # Özet
                if ai_analysis.get('summary'):
                    doc.add_heading('Metin Özeti', level=2)
                    doc.add_paragraph(str(ai_analysis.get('summary', '')))
                
                # Anahtar kelimeler
                if ai_analysis.get('keywords'):
                    doc.add_heading('Anahtar Kelimeler', level=2)
                    keywords_para = doc.add_paragraph()
                    for keyword in ai_analysis['keywords']:
                        keywords_para.add_run(f"• {keyword}\n")
                
                # Duygusal analiz
                if ai_analysis.get('emotion_analysis'):
                    doc.add_heading('Duygusal Analiz', level=2)
                    doc.add_paragraph(str(ai_analysis.get('emotion_analysis', '')))
            
            # Footer
            footer_section = doc.sections[0]
            footer = footer_section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Bu rapor Whisper AI tarafından {datetime.now().strftime('%d.%m.%Y %H:%M:%S')} tarihinde otomatik olarak oluşturulmuştur."
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Belgeyi buffer'a kaydet
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Word creation error: {e}")
            st.error(f"Word belgesi oluşturma hatası: {str(e)}")
            return None


class ExcelExporter:
    """Gelişmiş Excel export sınıfı"""
    
    @staticmethod
    def create_excel_report(transcription_data, ai_analysis=None):
        """Gelişmiş Excel raporu oluştur"""
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.utils.dataframe import dataframe_to_rows
            from io import BytesIO
            
            wb = Workbook()
            
            # Ana sayfa - Özet bilgiler
            ws_summary = wb.active
            ws_summary.title = "Özet"
            
            # Başlık
            ws_summary['A1'] = 'Whisper AI Transkripsiyon Raporu'
            ws_summary['A1'].font = Font(size=16, bold=True)
            ws_summary['A1'].alignment = Alignment(horizontal='center')
            ws_summary.merge_cells('A1:B1')
            
            # Rapor bilgileri
            summary_data = [
                ['Dosya Adı', str(transcription_data.get('file_name', 'Bilinmiyor'))],
                ['İşlem Tarihi', datetime.now().strftime("%d.%m.%Y %H:%M:%S")],
                ['Dil', str(transcription_data.get('language', 'Otomatik'))],
                ['İşlem Süresi (saniye)', f"{transcription_data.get('processing_time', 0):.1f}"],
                ['Dosya Boyutu (MB)', f"{transcription_data.get('file_size_mb', 0):.1f}"],
                ['ID', str(transcription_data.get('id', 'N/A'))]
            ]
            
            for i, (label, value) in enumerate(summary_data, start=3):
                ws_summary[f'A{i}'] = label
                ws_summary[f'B{i}'] = value
                ws_summary[f'A{i}'].font = Font(bold=True)
            
            # Transkripsiyon sayfası
            ws_transcript = wb.create_sheet("Transkripsiyon")
            
            transcript_text = transcription_data.get('transcript_text', 'Transkripsiyon metni bulunamadı.')
            
            ws_transcript['A1'] = 'Transkripsiyon Metni'
            ws_transcript['A1'].font = Font(size=14, bold=True)
            
            # Metin istatistikleri
            word_count = len(transcript_text.split())
            char_count = len(transcript_text)
            
            ws_transcript['A3'] = f'Kelime Sayısı: {word_count}'
            ws_transcript['A4'] = f'Karakter Sayısı: {char_count}'
            
            # Transkripsiyon metni - uzun metinleri parçalara böl
            max_cell_length = 32767  # Excel hücre sınırı
            if len(transcript_text) > max_cell_length:
                # Metni parçalara böl
                parts = []
                while transcript_text:
                    if len(transcript_text) <= max_cell_length:
                        parts.append(transcript_text)
                        break
                    else:
                        # En yakın paragraf sonunu bul
                        split_point = transcript_text.rfind('\n\n', 0, max_cell_length)
                        if split_point == -1:
                            split_point = max_cell_length
                        parts.append(transcript_text[:split_point])
                        transcript_text = transcript_text[split_point:].lstrip()
                
                for i, part in enumerate(parts):
                    ws_transcript[f'A{6+i}'] = part
                    ws_transcript[f'A{6+i}'].alignment = Alignment(wrap_text=True, vertical='top')
            else:
                ws_transcript['A6'] = transcript_text
                ws_transcript['A6'].alignment = Alignment(wrap_text=True, vertical='top')
            
            # AI analiz sayfası
            if ai_analysis and ai_analysis.get('summary'):
                ws_ai = wb.create_sheet("AI Analiz")
                
                ws_ai['A1'] = 'AI Analiz Sonuçları'
                ws_ai['A1'].font = Font(size=14, bold=True)
                
                row_num = 3
                
                # Özet
                if ai_analysis.get('summary'):
                    ws_ai[f'A{row_num}'] = 'Özet:'
                    ws_ai[f'A{row_num}'].font = Font(bold=True)
                    row_num += 1
                    
                    summary_text = str(ai_analysis.get('summary', ''))
                    if len(summary_text) > max_cell_length:
                        summary_text = summary_text[:max_cell_length-3] + "..."
                    
                    ws_ai[f'A{row_num}'] = summary_text
                    ws_ai[f'A{row_num}'].alignment = Alignment(wrap_text=True, vertical='top')
                    row_num += 3
                
                # Anahtar kelimeler
                if ai_analysis.get('keywords'):
                    ws_ai[f'A{row_num}'] = 'Anahtar Kelimeler:'
                    ws_ai[f'A{row_num}'].font = Font(bold=True)
                    row_num += 1
                    
                    for keyword in ai_analysis['keywords']:
                        ws_ai[f'A{row_num}'] = f'• {keyword}'
                        row_num += 1
                    
                    row_num += 2
                
                # Duygusal analiz
                if ai_analysis.get('emotion_analysis'):
                    ws_ai[f'A{row_num}'] = 'Duygusal Analiz:'
                    ws_ai[f'A{row_num}'].font = Font(bold=True)
                    row_num += 1
                    
                    emotion_text = str(ai_analysis.get('emotion_analysis', ''))
                    if len(emotion_text) > max_cell_length:
                        emotion_text = emotion_text[:max_cell_length-3] + "..."
                    
                    ws_ai[f'A{row_num}'] = emotion_text
                    ws_ai[f'A{row_num}'].alignment = Alignment(wrap_text=True, vertical='top')
            
            # Excel dosyasını buffer'a kaydet
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Excel creation error: {e}")
            st.error(f"Excel raporu oluşturma hatası: {str(e)}")
            return None


class QRCodeGenerator:
    """QR kod oluşturucu sınıfı"""
    
    @staticmethod
    def create_qr_code(text_content, file_id):
        """QR kod oluştur"""
        try:
            import qrcode
            from PIL import Image
            from io import BytesIO
            
            # QR kod oluştur
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            
            # Metin çok uzunsa kısalt
            if len(text_content) > 500:
                text_content = text_content[:500] + "..."
            
            qr.add_data(text_content)
            qr.make(fit=True)
            
            # QR kod görselini oluştur
            img = qr.make_image(fill_color="black", back_color="white")
            
            # Buffer'a kaydet
            buffer = BytesIO()
            img.save(buffer, format='PNG')
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"QR code creation error: {e}")
            st.error(f"QR kod oluşturma hatası: {str(e)}")
            return None


class ZipArchiver:
    """ZIP arşiv oluşturucu sınıfı"""
    
    @staticmethod
    def create_complete_archive(transcription_data, ai_analysis=None):
        """Komple ZIP arşivi oluştur"""
        try:
            import zipfile
            from io import BytesIO
            import json
            
            buffer = BytesIO()
            
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                file_name = transcription_data.get('file_name', 'transcript')
                base_name = os.path.splitext(file_name)[0]
                
                # Metin dosyası
                text_content = transcription_data.get('transcript_text', '')
                zf.writestr(f"{base_name}_transcript.txt", text_content)
                
                # JSON metadata
                metadata = {
                    'transcription': transcription_data,
                    'ai_analysis': ai_analysis,
                    'export_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'version': 'Whisper AI v2.0'
                }
                zf.writestr(f"{base_name}_metadata.json", json.dumps(metadata, indent=2, ensure_ascii=False))
                
                # README dosyası
                readme_content = f"""Whisper AI Transkripsiyon Arşivi
{'='*40}

Dosya: {file_name}
Tarih: {transcription_data.get('created_at', 'N/A')}
Dil: {transcription_data.get('language', 'Otomatik')}
İşlem Süresi: {transcription_data.get('processing_time', 0):.1f} saniye

İçerik:
- {base_name}_transcript.txt: Transkripsiyon metni
- {base_name}_metadata.json: AI analiz sonuçları ve metadata
- README.txt: Bu dosya

Whisper AI ile oluşturulmuştur.
"""
                zf.writestr("README.txt", readme_content)
                
                # AI analiz varsa ayrı dosya olarak ekle
                if ai_analysis and ai_analysis.get('summary'):
                    ai_content = f"""AI ANALİZ RAPORU
{'='*20}

Model: {ai_analysis.get('model', 'GPT-4-Turbo')}
Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}

ÖZET:
{ai_analysis.get('summary', 'Özet mevcut değil')}

ANAHTAR KELİMELER:
{', '.join(ai_analysis.get('keywords', []))}

DUYGUSAL ANALİZ:
{ai_analysis.get('emotion_analysis', 'Duygusal analiz mevcut değil')}
"""
                    zf.writestr(f"{base_name}_ai_analiz.txt", ai_content)
            
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"ZIP creation error: {e}")
            st.error(f"ZIP arşivi oluşturma hatası: {str(e)}")
            return None


class EmailSender:
    """Email gönderici sınıfı (placeholder)"""
    
    @staticmethod
    def send_transcription_email(recipient, transcription_data, ai_analysis=None, attachments=None):
        """Email gönderimi simülasyonu"""
        try:
            # Bu fonksiyon gerçek email entegrasyonu için hazır
            # SMTP ayarları ve email template'i burada olacak
            
            logger.info(f"Email would be sent to: {recipient}")
            return True, f"Email gönderimi simüle edildi: {recipient}"
            
        except Exception as e:
            logger.error(f"Email sending error: {e}")
            return False, f"Email gönderimi başarısız: {str(e)}"
